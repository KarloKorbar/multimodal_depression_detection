"""
Requirements:
- markdown2
- python-docx
- beautifulsoup4
- plantuml (optional, for PlantUML diagram rendering)
- pygments (for syntax highlighting)
- Java (for PlantUML rendering)
"""
import os
import markdown2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
import tempfile
import re

# Pygments for syntax highlighting
from pygments import highlight
from pygments.lexers import get_lexer_by_name, TextLexer
from pygments.formatter import Formatter
from pygments.token import Token
from pygments.styles import get_style_by_name
from pygments.formatters import ImageFormatter
from PIL import Image
import io
import textwrap

CHAPTERS = [
    "0_intro.md",
    "1_psychological_aspects.md",
    "2_mental_health_detection.md",
    "3_data_analysis_&_preprocessing.md",
    "4_implementing_individual_modalities.md",
    "5_implementing_multimodal_fusion.md",
    "6_evaluation.md",
]

CHAPTERS_DIR = "Chapters"
OUTPUT_DOCX = "Thesis_Combined.docx"

# Try to import plantuml for diagram rendering
try:
    from plantuml import PlantUML
    PLANTUML_SERVER = os.environ.get("PLANTUML_SERVER", "http://www.plantuml.com/plantuml/img/")
    plantuml = PlantUML(url=PLANTUML_SERVER)
    PLANTUML_AVAILABLE = True
except ImportError:
    PLANTUML_AVAILABLE = False

from bs4 import BeautifulSoup

def normalize_language(lang):
    if not lang:
        return None
    lang = lang.lower()
    aliases = {
        'py': 'python',
        'python3': 'python',
        'js': 'javascript',
        'ts': 'typescript',
        'c++': 'cpp',
        'sh': 'bash',
        'shell': 'bash',
        'console': 'bash',
        'text': 'text',
        'plaintext': 'text',
        'md': 'markdown',
        'yml': 'yaml',
        'ps1': 'powershell',
        'c#': 'csharp',
        'html': 'html',
        'xml': 'xml',
        'json': 'json',
        'css': 'css',
        'scss': 'scss',
        'go': 'go',
        'golang': 'go',
        'r': 'r',
        'matlab': 'matlab',
        'plantuml': 'plantuml',
    }
    return aliases.get(lang, lang)

def wrap_code_lines(code, width=80, subsequent_indent='    '):
    wrapped_lines = []
    for line in code.splitlines():
        if len(line) > width:
            wrapped = textwrap.fill(line, width=width, subsequent_indent=subsequent_indent, replace_whitespace=False)
            wrapped_lines.append(wrapped)
        else:
            wrapped_lines.append(line)
    return '\n'.join(wrapped_lines)

def add_code_block(doc, code, language=None):
    # Wrap code lines to a max width (e.g., 80 chars)
    code = wrap_code_lines(code, width=62)
    from pygments.lexers import get_lexer_by_name, guess_lexer, TextLexer
    lexer = None
    lang = normalize_language(language)
    if lang:
        try:
            lexer = get_lexer_by_name(lang, stripall=True)
        except Exception:
            lexer = None
    if lexer is None:
        try:
            lexer = guess_lexer(code)
        except Exception:
            lexer = TextLexer()
    # Use a dark theme for the image and a fixed font size for consistency
    formatter = ImageFormatter(style='monokai', font_name='Menlo', fontsize=14, line_numbers=False)
    # fontsize=14 ensures consistent, readable code block images
    img_data = highlight(code, lexer, formatter)
    img = Image.open(io.BytesIO(img_data))
    # Save to a temp file and insert into docx
    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
        img.save(tmp_img, format='PNG')
        tmp_img_path = tmp_img.name
    # Calculate height in inches for consistent font size
    dpi = img.info.get('dpi', (96, 96))[1]  # vertical DPI
    if dpi == 0: dpi = 96  # fallback
    height_inches = img.height / dpi
    # Insert with fixed height (e.g., 1.5 inches)
    doc.add_picture(tmp_img_path, height=Inches(height_inches))
    # Clean up temp file
    import os
    os.remove(tmp_img_path)
    return None

def add_plantuml_diagram(doc, plantuml_code):
    if not PLANTUML_AVAILABLE:
        doc.add_paragraph("[PlantUML diagram could not be rendered: plantuml package not installed]")
        add_code_block(doc, plantuml_code, language='plantuml')
        return
    with tempfile.NamedTemporaryFile(delete=False, suffix='.uml', mode='w', encoding='utf-8') as uml_file:
        uml_file.write(plantuml_code)
        uml_file_path = uml_file.name
    img_path = uml_file_path.replace('.uml', '.png')
    try:
        plantuml.processes_file(uml_file_path)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(5.5))
            os.remove(img_path)
        else:
            doc.add_paragraph("[PlantUML diagram could not be rendered: image not found]")
            add_code_block(doc, plantuml_code, language='plantuml')
    except Exception as e:
        doc.add_paragraph(f"[PlantUML diagram rendering error: {e}]")
        add_code_block(doc, plantuml_code, language='plantuml')
    finally:
        os.remove(uml_file_path)

def is_table_separator(line):
    # Checks if a line is a markdown table separator (e.g. |---|---|)
    return bool(re.match(r'\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$', line))

def parse_markdown_table(table_lines):
    # Parse markdown table lines into a list of rows (each row is a list of cells)
    rows = []
    for line in table_lines:
        # Remove leading/trailing whitespace and pipes, then split
        line = line.strip().strip('|')
        cells = [cell.strip() for cell in line.split('|')]
        rows.append(cells)
    return rows

def add_markdown_table(doc, table_lines):
    rows = parse_markdown_table([l for l in table_lines if not is_table_separator(l)])
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.cell(i, j).text = cell
    # Make first row bold (header)
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    return table

def add_markdown_to_docx(md_text, doc):
    lines = md_text.splitlines()
    in_code = False
    code_lines = []
    code_lang = None
    in_table = False
    table_lines = []
    for i, line in enumerate(lines):
        # Detect start of code block
        if not in_code and (line.strip().startswith('```') or line.strip().startswith('~~~')):
            in_code = True
            code_lang = line.strip().strip('`~').strip()
            if code_lang == '':
                code_lang = None
            code_lines = []
            continue
        # Detect end of code block
        if in_code and (line.strip().startswith('```') or line.strip().startswith('~~~')):
            code = '\n'.join(code_lines)
            if code.strip().startswith('@startuml') and code.strip().endswith('@enduml'):
                add_plantuml_diagram(doc, code)
            else:
                add_code_block(doc, code, language=code_lang)
            in_code = False
            code_lines = []
            code_lang = None
            continue
        if in_code:
            code_lines.append(line)
            continue
        # Table detection
        if not in_table and line.strip().startswith('|') and '|' in line.strip()[1:]:
            in_table = True
            table_lines = [line]
            continue
        if in_table:
            if line.strip().startswith('|') and '|' in line.strip()[1:]:
                table_lines.append(line)
                continue
            elif is_table_separator(line):
                table_lines.append(line)
                continue
            else:
                add_markdown_table(doc, table_lines)
                in_table = False
                table_lines = []
        # If we just finished a table, process this line as normal
        if not in_table and (line.strip() != ''):
            # Use markdown2/BeautifulSoup for headings, lists, etc.
            html_line = markdown2.markdown(line, extras=["tables"])
            soup_line = BeautifulSoup(html_line, "html.parser")
            for elem in soup_line.children:
                if elem.name and elem.name.startswith('h') and elem.name[1:].isdigit():
                    level = int(elem.name[1:])
                    p = doc.add_paragraph(elem.get_text(), style=f'Heading {min(level, 4)}')
                elif elem.name == 'p':
                    doc.add_paragraph(elem.get_text())
                elif elem.name == 'ul':
                    for li in elem.find_all('li', recursive=False):
                        p = doc.add_paragraph(li.get_text(), style='List Bullet')
                elif elem.name == 'ol':
                    for li in elem.find_all('li', recursive=False):
                        p = doc.add_paragraph(li.get_text(), style='List Number')
                elif elem.name == 'blockquote':
                    p = doc.add_paragraph(elem.get_text())
                    p.style = 'Intense Quote'
                elif elem.name == 'hr':
                    doc.add_page_break()
    # If file ends with a table
    if in_table and table_lines:
        add_markdown_table(doc, table_lines)

def main():
    doc = Document()
    for chapter in CHAPTERS:
        chapter_path = os.path.join(CHAPTERS_DIR, chapter)
        if not os.path.exists(chapter_path):
            print(f"Warning: {chapter_path} not found.")
            continue
        with open(chapter_path, encoding='utf-8') as f:
            md_text = f.read()
        add_markdown_to_docx(md_text, doc)
        doc.add_page_break()
    doc.save(OUTPUT_DOCX)
    print(f"Combined .docx saved as {OUTPUT_DOCX}")

if __name__ == "__main__":
    main() 